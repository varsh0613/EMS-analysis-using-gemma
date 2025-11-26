# doc_builder.py
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

class DocBuilder:
    """
    Build Word report for EDA: summary, tables, plots, interpretations, insights.
    """

    def __init__(self, save_dir=None):
        if save_dir is None:
            self.save_dir = os.path.join(os.path.expanduser("~"), "Desktop", "uni", "gemma")
        else:
            self.save_dir = save_dir
        os.makedirs(self.save_dir, exist_ok=True)

    def build_report(self, summary_text, df, plot_paths, interpretations, final_insights=None, recommendations=None, filename="eda_report.docx"):
        doc = Document()
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)

        doc.add_heading("Exploratory Data Analysis Report", 0)

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(summary_text)

        # Dataset Snapshot
        doc.add_heading("Dataset Snapshot", level=1)
        doc.add_paragraph(f"Rows: {df.shape[0]}, Columns: {df.shape[1]}")
        numeric_cols = df.select_dtypes(include=['int64','float64']).columns
        cat_cols = df.select_dtypes(include=['object','category']).columns
        doc.add_paragraph(f"Numeric columns: {len(numeric_cols)} ({', '.join(numeric_cols)})")
        doc.add_paragraph(f"Categorical columns: {len(cat_cols)} ({', '.join(cat_cols)})")

        # Missing values
        doc.add_heading("Missing Values", level=2)
        missing = df.isna().sum()
        missing = missing[missing > 0].sort_values(ascending=False)
        if not missing.empty:
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Column"
            hdr_cells[1].text = "Missing Count"
            for col, cnt in missing.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(col)
                row_cells[1].text = str(cnt)
        else:
            doc.add_paragraph("No missing values detected.")

        # Numeric summary
        if len(numeric_cols) > 0:
            doc.add_heading("Numeric Summary Statistics", level=1)
            stats = df[numeric_cols].describe().transpose()
            table = doc.add_table(rows=1, cols=len(stats.columns)+1)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Column"
            for i, col in enumerate(stats.columns):
                hdr_cells[i+1].text = str(col)
            for idx, row in stats.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                for i, val in enumerate(row):
                    row_cells[i+1].text = f"{val:.3f}"

        # Plots and interpretations
        doc.add_heading("Plots & Interpretations", level=1)
        for pf in plot_paths:
            fp = os.path.join(self.save_dir, pf)
            if os.path.exists(fp):
                try:
                    doc.add_picture(fp, width=Inches(6))
                except:
                    continue
                interp = interpretations.get(pf, "No interpretation available.")
                doc.add_paragraph(interp)

        # Final Insights
        doc.add_heading("Final Insights", level=1)
        if final_insights:
            for ins in final_insights:
                doc.add_paragraph("- " + ins)
        # Recommendations
        doc.add_heading("Recommendations", level=1)
        if recommendations:
            for rec in recommendations:
                doc.add_paragraph("- " + rec)

        save_path = os.path.join(self.save_dir, filename)
        doc.save(save_path)
        return save_path
